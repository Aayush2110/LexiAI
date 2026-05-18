"""
Document Loader Service

Handles loading and extracting text from various document formats.

Supported Formats:
- PDF: Using PyMuPDF (fitz) for better text extraction
- DOCX: Using python-docx for Word documents
- TXT: Plain text files

Why Multiple Loaders?
Different document formats require different parsing libraries.
PyMuPDF is faster and more accurate than PyPDF2 for PDFs.
"""

from typing import List
import os
from loguru import logger
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document
from app.utils.helpers import clean_text


class DocumentLoader:
    """
    Document Loader Class
    
    Loads and extracts text from various document formats.
    Returns structured data with metadata for better tracking.
    """
    
    def __init__(self):
        """Initialize document loader"""
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """
        Load multiple documents and extract text
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List[Document]: List of LangChain Document objects with text and metadata
        """
        
        documents = []
        
        for file_path in file_paths:
            try:
                logger.info(f"Loading document: {file_path}")
                
                # Determine file format
                extension = file_path.split('.')[-1].lower()
                
                if extension == 'pdf':
                    docs = self._load_pdf(file_path)
                elif extension == 'docx':
                    docs = self._load_docx(file_path)
                elif extension == 'txt':
                    docs = self._load_txt(file_path)
                else:
                    logger.warning(f"Unsupported format: {extension}")
                    continue
                
                documents.extend(docs)
                logger.info(f"Loaded {len(docs)} pages/sections from {os.path.basename(file_path)}")
            
            except Exception as e:
                logger.error(f"Error loading {file_path}: {str(e)}")
                continue
        
        return documents
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """
        Load PDF file using PyMuPDF
        
        PyMuPDF (fitz) is faster and more accurate than PyPDF2.
        It handles complex PDFs with images and tables better.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List[Document]: LangChain Documents with page-level granularity
        """
        
        documents = []
        
        try:
            # Open PDF
            pdf_document = fitz.open(file_path)
            
            # Extract text from each page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                
                # Clean and validate text
                text = clean_text(text)
                
                if text.strip():  # Only add non-empty pages
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            'source': os.path.basename(file_path),
                            'page': page_num + 1,
                            'format': 'pdf',
                            'total_pages': len(pdf_document)
                        }
                    ))
            
            pdf_document.close()
        
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """
        Load DOCX file using python-docx
        
        Extracts text from Word documents paragraph by paragraph.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List[Document]: LangChain Document with combined text
        """
        
        documents = []
        
        try:
            # Open DOCX
            doc = DocxDocument(file_path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Combine into single text
            text = '\n'.join(paragraphs)
            text = clean_text(text)
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        'source': os.path.basename(file_path),
                        'page': 1,
                        'format': 'docx',
                        'total_paragraphs': len(paragraphs)
                    }
                ))
        
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
        
        return documents
    
    def _load_txt(self, file_path: str) -> List[Document]:
        """
        Load plain text file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            List[Document]: LangChain Document with text content
        """
        
        documents = []
        
        try:
            # Read text file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            text = clean_text(text)
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        'source': os.path.basename(file_path),
                        'page': 1,
                        'format': 'txt'
                    }
                ))
        
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            raise
        
        return documents


# Global instance
document_loader = DocumentLoader()
